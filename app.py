# app.py

import streamlit as st
import os
import re
import psutil
import pyttsx3
import asyncio
from datetime import datetime
from docx import Document
from config import OUTPUT_DIR

from agent_core import ContentAgent
from local_llm_loader import load_local_llm
from data_visualizer import DataVisualizer
# --- THIS IS THE FIX: Import from the newly named file ---
from ollama_monitor import OllamaMonitor 

def text_to_speech(text: str):
    try:
        engine = pyttsx3.init(); engine.say(text); engine.runAndWait()
    except Exception as e:
        st.error(f"Could not initialize Text-to-Speech engine: {e}")

def save_document(filename: str, title: str, final_content: list, visual_mapping: dict):
    if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
    doc = Document(); doc.add_heading(title, level=1)
    for section in final_content:
        doc.add_heading(section['topic'], level=2)
        parts = re.split(r'(<FIGURE>.*?</FIGURE>)', section.get('content', ''), flags=re.DOTALL)
        for part in parts:
            if not part.strip(): continue
            if part.startswith('<FIGURE>'):
                visual_match = re.search(r'(\[CHART\|.*?\])', part, re.DOTALL)
                caption_match = re.search(r'<CAPTION>(.*?)</CAPTION>', part, re.DOTALL)
                if visual_match:
                    placeholder = visual_match.group(1)
                    if placeholder in visual_mapping and visual_mapping.get(placeholder):
                        try:
                            doc.add_picture(visual_mapping[placeholder], width=Inches(6.0))
                            if caption_match: doc.add_paragraph(caption_match.group(1).strip(), style='Caption')
                        except Exception as e: doc.add_paragraph(f"[Error adding chart: {e}]", style='Comment')
                    else: doc.add_paragraph(f"[Chart generation failed for: {placeholder}]", style='Comment')
            else:
                if part.strip().startswith('|') and '|' in part:
                    try:
                        lines = [l.strip() for l in part.strip().split('\n')]; header = [h.strip() for h in lines[0].strip('|').split('|')]
                        table = doc.add_table(rows=1, cols=len(header)); table.style = 'Table Grid'
                        for i, h in enumerate(header): table.rows[0].cells[i].text = h
                        for line in lines[2:]:
                            row_cells = table.add_row().cells
                            for i, cell in enumerate(line.strip('|').split('|')): row_cells[i].text = cell.strip()
                    except Exception as e: doc.add_paragraph(f"[Error rendering table: {e}]\n{part}")
                else: doc.add_paragraph(part)
    sanitized_filename = filename.replace(" ", "_").lower()
    filepath = os.path.join(OUTPUT_DIR, f"{sanitized_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filepath); return filepath

st.set_page_config(layout="wide"); st.title("🎓 On-Device Agentic Researcher (Ollama Edition)")
for key in ['agent', 'chart_agent', 'document_sections', 'outline', 'doc_title', 'metrics', 'processing', 'summary_section']:
    if key not in st.session_state:
        defaults = {'processing': False, 'document_sections': [], 'outline': "", 'doc_title': "", 'summary_section': None}
        st.session_state[key] = defaults.get(key, None)

with st.sidebar:
    st.header("⚙️ Configuration")
    st.info("Ensure the Ollama application is running and the model has been pulled (e.g., `ollama pull phi3`).")
    if st.button("Connect to Ollama", disabled=st.session_state.processing):
        st.session_state.processing = True; st.rerun() 

if st.session_state.processing and st.session_state.agent is None:
    try:
        with st.spinner("Connecting to Ollama server..."):
            st.session_state.llm = load_local_llm()
            st.session_state.agent = ContentAgent(st.session_state.llm)
        st.success("Connected to Ollama & Agent is Ready!")
        st.session_state.chart_agent = DataVisualizer()
        st.success("Data Visualizer Ready!")
    except Exception as e:
        st.error(f"A critical error occurred while connecting to Ollama: {e}", icon="🚨")
    finally:
        st.session_state.processing = False; st.rerun()

if st.session_state.agent:
    st.header("1. Define Your Research Document")
    doc_title_input = st.text_input("Document Title:", "The Future of Artificial Intelligence")
    user_request = st.text_area("Your Request/Abstract:", "Write a comprehensive paper on the future of AI.", height=150)
    
    if st.button("✍️ Generate Initial Draft", disabled=st.session_state.processing):
        st.session_state.doc_title = doc_title_input
        st.session_state.processing = True
        st.session_state.document_sections = []; st.session_state.summary_section = None
        
        with st.spinner("The agent is planning and drafting concurrently..."):
            monitor = None
            try:
                monitor = OllamaMonitor()
                monitor.start()
                
                outline, sections, metrics = asyncio.run(
                    st.session_state.agent.run_initial_draft_async(user_request, st.session_state.doc_title)
                )
                
                avg_cpu, avg_mem, avg_gpu = monitor.stop()
                metrics["avg_cpu_utilization"] = avg_cpu
                metrics["avg_memory_utilization"] = avg_mem
                metrics["avg_gpu_utilization"] = avg_gpu

                st.session_state.outline = outline; st.session_state.document_sections = sections; st.session_state.metrics = metrics
                
                if sections:
                    st.session_state.summary_section = sections[0]
                
            except Exception as e:
                if monitor: monitor.stop()
                st.error(f"An error occurred during document generation: {e}")
            finally:
                st.session_state.processing = False
                st.rerun()
else:
    if not st.session_state.processing:
        st.warning("Please connect to the Ollama server using the sidebar.")

if st.session_state.metrics:
    st.header("📊 Advanced Performance Metrics")
    m = st.session_state.metrics
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Generation Time", f"{m.get('total_time', 0):.2f} s")
    col2.metric("Total Tokens Generated", f"{m.get('total_tokens', 0)}")
    col3.metric("Time to First Token (TTFT)", f"{m.get('ttft', 0) * 1000:.0f} ms", help="Time until the first piece of content (the outline) was generated.")
    #col4.metric("Inter-Token Latency", f"{m.get('inter_token_latency', 0):.2f} ms/token")
    # We now display Tokens per Second, which is a more intuitive metric.
    col4.metric("Tokens per Second (t/s)", f"{m.get('tokens_per_second', 0):.2f} t/s", help="The overall throughput of the content generation phase.")

    with st.expander("View System Utilization (Averages for Ollama process)"):
        c1, c2, c3 = st.columns(3)
        c1.metric("Average CPU Utilization", f"{m.get('avg_cpu_utilization', 0):.1f}%")
        c2.metric("Average Memory Utilization", f"{m.get('avg_memory_utilization', 0):.1f}%")
        c3.metric("Average GPU Utilization", f"{m.get('avg_gpu_utilization', 0):.1f}%")

if st.session_state.document_sections:
    st.header("2. Review and Refine Draft")
    review_sections = st.session_state.document_sections[1:] if st.session_state.summary_section else st.session_state.document_sections
    for i, section in enumerate(review_sections):
        with st.container(border=True):
            st.subheader(f"Section: {section['topic']}")
            st.markdown(section['content'])
            feedback = st.text_area("Your feedback for this section:", key=f"feedback_{i}")
            if st.button("🔄 Regenerate Section", key=f"regen_{i}", disabled=st.session_state.processing):
                st.session_state.processing = True
                with st.spinner("The agent is revising this section..."):
                    original_index = i + 1 if st.session_state.summary_section else i
                    revised_content = st.session_state.agent.regenerate_section(st.session_state.doc_title, st.session_state.outline.split('\n'), section['topic'], section['content'], feedback)
                    st.session_state.document_sections[original_index]['content'] = revised_content
                st.session_state.processing = False; st.rerun()

if st.session_state.summary_section:
    st.header("3. Document Summary & Audio Transcript")
    with st.container(border=True):
        st.subheader(st.session_state.summary_section['topic'])
        st.markdown(st.session_state.summary_section['content'])
        
        if st.button("🔊 Generate Audio Transcript", disabled=st.session_state.processing):
            st.session_state.processing = True
            with st.spinner("Generating audio file..."):
                from audio_utils import generate_audio_file
                transcript = f"Title: {st.session_state.doc_title}. {st.session_state.summary_section['topic']}: {st.session_state.summary_section['content']}"
                audio_path = generate_audio_file(transcript, st.session_state.doc_title.replace(" ", "_").lower())
                st.session_state.audio_file_path = audio_path
            st.session_state.processing = False
            st.rerun()

        if 'audio_file_path' in st.session_state and st.session_state.audio_file_path:
            st.audio(st.session_state.audio_file_path)

if st.session_state.document_sections:
    st.header("4. Finalize and Build Document")
    if st.button("✅ Build .docx File", type="primary", disabled=st.session_state.processing):
        with st.spinner("Assembling .docx file..."):
            visual_mapping = {}
            st.write("Parsing document for charts...")
            for section in st.session_state.document_sections:
                figures = re.findall(r'(<FIGURE>.*?</FIGURE>)', section['content'], flags=re.DOTALL)
                for block in figures:
                    chart_match = re.search(r'(\[CHART\|(bar|pie):\s*({.*})\])', block, re.DOTALL)
                    if chart_match:
                        full_placeholder, chart_type, chart_json = chart_match.groups()
                        fname = f"{st.session_state.doc_title.replace(' ', '_').lower()}_visual_{len(visual_mapping)+1}"
                        path = st.session_state.chart_agent.generate_chart(chart_type, chart_json, fname)
                        visual_mapping[full_placeholder] = path
            filepath = save_document(st.session_state.doc_title, st.session_state.doc_title, st.session_state.document_sections, visual_mapping)
            st.success(f"Document built successfully! ✨"); st.markdown(f"**Download from:** `{filepath}`")